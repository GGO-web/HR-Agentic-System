import os
from typing import List, Dict, Any, Optional
from pathlib import Path
import pdfplumber
from docx import Document as DocxDocument
from langchain_core.documents import Document


class DocumentLoader:
    """Class for loading and parsing documents (PDF and DOCX)."""

    @staticmethod
    async def load_document(file_path: str, metadata: Optional[Dict[str, Any]] = None) -> List[Document]:
        """
        Load document from file path and extract text with metadata.

        Args:
            file_path: Path to the document file
            metadata: Additional metadata to attach to the document

        Returns:
            List of LangChain Document objects with extracted text and metadata
        """
        file_path_obj = Path(file_path)

        if not file_path_obj.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        file_extension = file_path_obj.suffix.lower()
        base_metadata = {
            "filename": file_path_obj.name,
            "source_type": file_extension[1:] if file_extension else "unknown",
            "file_path": str(file_path_obj),
        }

        if metadata:
            base_metadata.update(metadata)

        if file_extension == ".pdf":
            text = await DocumentLoader._load_pdf(file_path)
        elif file_extension in [".docx", ".doc"]:
            text = await DocumentLoader._load_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")

        return [Document(page_content=text, metadata=base_metadata)]

    @staticmethod
    async def _load_pdf(file_path: str) -> str:
        """
        Extract text from PDF file using pdfplumber.
        Preserves structure and avoids mixing columns.

        Args:
            file_path: Path to PDF file

        Returns:
            Extracted text content
        """
        text_parts = []

        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                # Try multiple extraction methods for better text recovery
                # Method 1: Standard text extraction
                page_text = page.extract_text()

                # Method 2: If standard extraction is poor, try extracting from words
                if not page_text or len(page_text.strip()) < 50:
                    words = page.extract_words()
                    if words:
                        page_text = ' '.join(
                            [word.get('text', '') for word in words if word.get('text')])

                if page_text and page_text.strip():
                    text_parts.append(page_text)

        return "\n\n".join(text_parts)

    @staticmethod
    async def _load_docx(file_path: str) -> str:
        """
        Extract text from DOCX file using python-docx.

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text content
        """
        doc = DocxDocument(file_path)
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        return "\n\n".join(text_parts)
